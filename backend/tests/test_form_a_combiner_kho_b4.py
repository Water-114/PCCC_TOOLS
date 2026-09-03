"""Form A gốc (A14/A15) — test dựng lại ĐÚNG kịch bản thực tế Kho B4 (Cty
Liên Anh) đã biết, theo yêu cầu Phần 4 của prompt "Form A gốc (A14 + A15) —
combiner, không dùng AI":

- Quy mô: kho 12.901 m² sàn, 151.788 m³, hạng C.
- pham_vi_de_nghi KHÔNG gồm tram_bom/dienpccc (đợt cải tạo chỉ xin thẩm định
  kệ cao tầng + họng nước + sprinkler + phương tiện ban đầu).
- ha_tang_hien_huu: 2 bản ghi (trạm bơm, điện PCCC) — GCN gốc 490/TD-PCCC
  (2015), cải tạo bổ sung 621/TD-PCCC-P2 (2016), nghiệm thu 273/CSPCCC-P2
  (2016) — đúng lỗi thực tế owner đã gặp (quy-tac-dien-form.md mục 4c).
- b_form_results: giakehang (B15, có KN) / hong_nuoc (B5, có KN) / den_su_co
  (B13, có KN) / binh_chua_chay (B12, có KN).

Build ra file A14 (mục "8" điện PCCC CHỈ tồn tại ở A14, không có ở A15 — dùng
A14 đúng như prompt yêu cầu để phủ luôn nhánh tail A14-only). Assert theo
ĐÚNG lưu ý của owner: KHÔNG so khớp nguyên văn, chỉ so Kết luận + đúng nhánh."""

import pytest

from app.services import credits, form_a_combiner, pham_vi_hien_huu_store, quy_mo_store


@pytest.fixture
def kho_b4_session(app, client):
    client.post("/api/auth/register", json={"email": "khob4@pccc.local", "password": "matkhau123"})
    resp = client.post("/api/auth/login", json={"email": "khob4@pccc.local", "password": "matkhau123"})
    data = resp.get_json()
    credits.grant_credits(data["user"]["id"], 5, credits.CREDIT_REASON_EMAIL_VERIFICATION, note="test")
    resp2 = client.post("/api/aiho/session/open", headers={"Authorization": f"Bearer {data['token']}"})
    return resp2.get_json()["session_id"]


_EMPTY_KIEN_NGHI = {"I_chua_the_hien": [], "II_chua_thong_nhat": [], "III_chua_phu_hop": [], "IV_de_xuat_bo_sung": []}


def _b_form_with_kn(so_hieu):
    """1 B-form co it nhat 1 item KN (chua_dat) - dung lai dung {id, noi_dung_thiet_ke,
    ket_luan} shape that cac reader tra ve, khong can du toan bo tieu chi that."""
    return {
        "items": [
            {"id": 1, "noi_dung_thiet_ke": "Đáp ứng theo bản vẽ (fake).", "ket_luan": "dat"},
            {"id": 2, "noi_dung_thiet_ke": "Chưa thể hiện rõ trên bản vẽ.", "ket_luan": "chua_the_hien"},
        ],
        "tong_ket": "Fake test.",
        "kien_nghi": _EMPTY_KIEN_NGHI,
        "so_hieu_ban_ve": so_hieu,
    }


@pytest.fixture
def kho_b4_setup(kho_b4_session):
    session_id = kho_b4_session
    quy_mo = {"occ": "kho", "totalArea": 12901, "volume": 151788, "hazard": "C", "floors": 1}
    quy_mo_store.save_quy_mo(session_id, quy_mo, source="manual")

    pham_vi_hien_huu_store.save_pham_vi_de_nghi(session_id, [
        "giakehang", "hong_nuoc", "chua_chay_tu_dong", "densucco", "binhchuachay",
        "baochay", "khibotsolkhi", "botcodinh", "botchuachay",
    ])  # KHONG gom tram_bom/dienpccc

    for ten_he_thong in ("tram_bom", "dienpccc"):
        pham_vi_hien_huu_store.save_ha_tang_hien_huu(
            session_id, ten_he_thong,
            gcn_so="490/TD-PCCC", gcn_ngay="15/01/2015",
            gcn_bo_sung_so="621/TD-PCCC-P2", gcn_bo_sung_ngay="20/06/2016",
            nghiem_thu_so="273/CSPCCC-P2", nghiem_thu_ngay="15/09/2016",
        )

    b_form_results = {
        "chua_chay_gia_ke_hang": _b_form_with_kn("B15-KHOB4"),
        "hong_nuoc": _b_form_with_kn("B5-KHOB4"),
        "den_su_co": _b_form_with_kn("B13-KHOB4"),
        "binh_chua_chay": _b_form_with_kn("B12-KHOB4"),
    }

    session_data = {"session_id": session_id, "quy_mo": quy_mo, "b_form_results": b_form_results}
    return session_data


def _row(doc, row_id):
    row = doc.tables[0].rows[row_id]
    return row.cells[2].text.strip(), row.cells[5].text.strip()


def test_kho_b4_tram_bom_hien_huu(kho_b4_setup):
    """Muc 3.2.4 (tram bom) = id 34 (A14) - Ket luan RONG, noi dung co 'hien
    huu' + dung so GCN bo sung 621/TD-PCCC-P2 (dung lo i thuc te owner da gap)."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    noi_dung, ket_luan = _row(doc, 34)
    assert ket_luan == ""
    assert "hiện hữu" in noi_dung
    assert "621/TD-PCCC-P2" in noi_dung


def test_kho_b4_dien_pccc_hien_huu(kho_b4_setup):
    """Muc 8 (dien PCCC) = id 63 (CHI co trong A14) - Ket luan RONG, noi dung
    co 'hien huu' + dung so GCN bo sung 621/TD-PCCC-P2."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    noi_dung, ket_luan = _row(doc, 63)
    assert ket_luan == ""
    assert "hiện hữu" in noi_dung
    assert "621/TD-PCCC-P2" in noi_dung


def test_kho_b4_cap_nuoc_ngoai_nha_khong_thuoc_dien(kho_b4_setup):
    """Muc 3.2.3 (cap nuoc ngoai nha) = id 32 - occ='kho' khong thuoc danh
    muc Phu luc C (evaluate_ngoai_nha tra 'na') -> Ket luan RONG, noi dung
    'Khong ap dung'."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    noi_dung, ket_luan = _row(doc, 32)
    assert ket_luan == ""
    assert "Không áp dụng" in noi_dung


def test_kho_b4_loa_khong_thuoc_dien(kho_b4_setup):
    """Muc 4.2 (loa) dong 'Doi tuong trang bi' = id 45 - occ='kho' khong
    thuoc danh sach TT1-6 Bang G.1 (evaluate_loa tra 'no') -> Ket luan RONG."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    _noi_dung, ket_luan = _row(doc, 45)
    assert ket_luan == ""


@pytest.mark.parametrize("row_id,label", [
    (25, "sprinkler (B15 - gia ke hang)"),
    (28, "hong nuoc (B5)"),
    (43, "den (B13)"),
    (53, "binh (B12)"),
])
def test_kho_b4_dan_chieu_b_form_co_kn(kho_b4_setup, row_id, label):
    """4 dong dan chieu B-form deu phai la KN vi b_form_results gia lap co
    it nhat 1 item chua_the_hien (khong phai dat/khong_ap_dung)."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    noi_dung, ket_luan = _row(doc, row_id)
    assert ket_luan == "KN", f"{label}: expected KN, got {ket_luan!r} ({noi_dung!r})"
    assert "kiến nghị" in noi_dung.lower()


def test_kho_b4_sprinkler_dan_chieu_finds_b15_not_b6(kho_b4_setup):
    """Rieng dong sprinkler (id=25): b_form_results CHI co 'chua_chay_gia_ke_hang'
    (B15), KHONG co 'chua_chay_tu_dong' (B6) - phai tim thay B15 (thu tu tim
    kiem trong _MUC['sprinkler'].b_forms: B6 truoc, B15 sau - B6 khong co du
    lieu nen roi xuong B15)."""
    from docx import Document
    import io
    docx_bytes = form_a_combiner.build_form_a_goc("A14", kho_b4_setup)
    doc = Document(io.BytesIO(docx_bytes))
    noi_dung, _ket_luan = _row(doc, 25)
    assert "B15" in noi_dung
