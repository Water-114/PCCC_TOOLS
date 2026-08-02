"""Test truc tiep mdc_filler.fill_docx() — to do cot "Ket luan" CHI khi gia
tri la "KN", giu mau mac dinh cho "Dat" va rong (muc "khong_ap_dung" duoc de
trong o tang route/aiho.py). Ap dung CHUNG cho moi hang muc (khong rieng
B12/B13) vi sua dung 1 cho trong ham dung chung."""

import io

from docx import Document

from app.services import mdc_filler


def _fill_and_read_ket_luan(loai, ket_luan_text, noi_dung="abc"):
    rows = mdc_filler.load_criteria_rows(loai)
    row_id = rows[0]["id"]
    docx_bytes = mdc_filler.fill_docx(loai, [{"id": row_id, "noi_dung_thiet_ke": noi_dung, "ket_luan": ket_luan_text}])
    doc = Document(io.BytesIO(docx_bytes))
    cell = doc.tables[0].rows[row_id].cells[mdc_filler.COL_KET_LUAN]
    runs = cell.paragraphs[0].runs
    color = runs[0].font.color.rgb if runs else None
    return cell.text, color


def test_ket_luan_kn_is_colored_red():
    text, color = _fill_and_read_ket_luan("dien_pccc", "KN")
    assert text == "KN"
    assert color == mdc_filler.MAU_DO


def test_ket_luan_dat_is_not_colored_red():
    text, color = _fill_and_read_ket_luan("dien_pccc", "Đạt")
    assert text == "Đạt"
    assert color != mdc_filler.MAU_DO


def test_ket_luan_empty_is_not_colored_red():
    """Mo phong "khong_ap_dung" (routes/aiho.py da doi thanh chuoi rong truoc
    khi goi fill_docx) - o rong, khong to do."""
    text, color = _fill_and_read_ket_luan("dien_pccc", "")
    assert text == ""
    assert color != mdc_filler.MAU_DO


def test_noi_dung_thiet_ke_still_colored_red_regardless_of_ket_luan():
    """Xac nhan hanh vi cu (cot 3, COL_THIET_KE) khong bi anh huong boi thay
    doi o cot 5 (COL_KET_LUAN)."""
    rows = mdc_filler.load_criteria_rows("dien_pccc")
    row_id = rows[0]["id"]
    docx_bytes = mdc_filler.fill_docx("dien_pccc", [{"id": row_id, "noi_dung_thiet_ke": "noi dung moi", "ket_luan": "Đạt"}])
    doc = Document(io.BytesIO(docx_bytes))
    cell = doc.tables[0].rows[row_id].cells[mdc_filler.COL_THIET_KE]
    assert cell.paragraphs[0].runs[0].font.color.rgb == mdc_filler.MAU_DO


def test_ket_luan_coloring_applies_to_other_forms_too():
    """Khong rieng 1 form nao - test them B12 (binh_chua_chay) de xac nhan
    dung chung cho moi hang muc, dung nhu chi dao (khong sua rieng tung
    reader)."""
    text, color = _fill_and_read_ket_luan("binh_chua_chay", "KN")
    assert text == "KN"
    assert color == mdc_filler.MAU_DO
