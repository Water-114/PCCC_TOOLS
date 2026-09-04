"""Batch 5A Pha 3 Buoc 3 — test route POST /api/aiho/export-cong-van-huong-dan:
nhan JSON hang_muc da co san (KHONG goi AI, KHONG tru quota), doc quy_mo tu DB
theo session_id, build .docx that tu file mau tra ve base64, van can dang nhap
+ dung chu so huu phien."""

import base64
import io

from docx import Document

from app.models import AIHO_API_NAME, count_usage_today
from app.services import credits, quy_mo_store


def _register_login_and_grant(client, email="cvhd1@pccc.local", password="matkhau123", amount=5):
    client.post("/api/auth/register", json={"email": email, "password": password})
    resp = client.post("/api/auth/login", json={"email": email, "password": password})
    data = resp.get_json()
    token, user_id = data["token"], data["user"]["id"]
    if amount:
        credits.grant_credits(user_id, amount, credits.CREDIT_REASON_EMAIL_VERIFICATION, note="test setup")
    return token, user_id


def _open_session(client, token):
    resp = client.post("/api/aiho/session/open", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 200, resp.get_json()
    return resp.get_json()["session_id"]


def _valid_hang_muc():
    return [
        {
            "slot": "baochay",
            "ten_he_thong": "Báo cháy tự động",
            "so_hieu_ban_ve": "BC-01",
            "kien_nghi": {
                "I_chua_the_hien": ["Thể hiện rõ ... (Điều 1)."],
                "II_chua_thong_nhat": [],
                "III_chua_phu_hop": [],
                "IV_de_xuat_bo_sung": [],
            },
        }
    ]


def test_requires_login(client):
    resp = client.post("/api/aiho/export-cong-van-huong-dan", json={"session_id": 1, "hang_muc": _valid_hang_muc()})
    assert resp.status_code == 401


def test_missing_session_id_returns_400(client):
    token, _ = _register_login_and_grant(client, email="cvhd2@pccc.local")
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"hang_muc": _valid_hang_muc()},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 400


def test_missing_hang_muc_returns_400(client):
    token, _ = _register_login_and_grant(client, email="cvhd3@pccc.local")
    session_id = _open_session(client, token)
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"session_id": session_id},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 400


def test_non_dict_root_returns_400(client):
    token, _ = _register_login_and_grant(client, email="cvhd4@pccc.local")
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json=[1, 2, 3],
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 400


def test_other_users_session_returns_404(client):
    token1, _ = _register_login_and_grant(client, email="cvhd5a@pccc.local")
    session_id = _open_session(client, token1)

    token2, _ = _register_login_and_grant(client, email="cvhd5b@pccc.local")
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"session_id": session_id, "hang_muc": _valid_hang_muc()},
        headers={"Authorization": f"Bearer {token2}"},
    )
    assert resp.status_code == 404


def test_valid_payload_returns_200_with_downloadable_docx(client):
    token, _ = _register_login_and_grant(client, email="cvhd6@pccc.local")
    session_id = _open_session(client, token)
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"session_id": session_id, "hang_muc": _valid_hang_muc()},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 200, resp.get_json()
    data = resp.get_json()
    assert data["filename"].endswith(".docx")
    docx_bytes = base64.b64decode(data["base64"])
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs]
    assert "- Thể hiện rõ ... (Điều 1)." in texts


def test_uses_saved_quy_mo_for_mergefields(client):
    """quy_mo da luu qua quy_mo_store.save_quy_mo() (vd tu /quymo-manual hoac
    /read-quymo truoc do trong CUNG phien) phai duoc doc lai va dien dung vao
    MERGEFIELD - route KHONG can nhan lai quy_mo tu frontend."""
    token, _ = _register_login_and_grant(client, email="cvhd7@pccc.local")
    session_id = _open_session(client, token)
    quy_mo_store.save_quy_mo(session_id, {"occ": "khachsan", "tenCongTrinh": "Test Building XYZ"}, source="manual")

    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"session_id": session_id, "hang_muc": _valid_hang_muc()},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 200, resp.get_json()
    docx_bytes = base64.b64decode(resp.get_json()["base64"])
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + "\n".join(p.text for p in cell.paragraphs)
    assert "Test Building XYZ" in full_text


def test_export_does_not_consume_aiho_quota(client):
    token, user_id = _register_login_and_grant(client, email="cvhd8@pccc.local")
    session_id = _open_session(client, token)
    before = count_usage_today(user_id, AIHO_API_NAME)
    resp = client.post(
        "/api/aiho/export-cong-van-huong-dan",
        json={"session_id": session_id, "hang_muc": _valid_hang_muc()},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert resp.status_code == 200
    after = count_usage_today(user_id, AIHO_API_NAME)
    assert after == before == 0
