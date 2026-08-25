"""B16 (chữa cháy bằng bột, TCVN 13877-2:2023) — khoá lại nội dung system
prompt của bot_chua_chay_reader.py: đúng 33 id (khớp danh sách owner đã trích
sẵn từ B16_bot_chua_chay.docx SAU KHI xoá bảng thừa 1x2, tự kiểm chứng lại qua
mdc_filler._extract_rows() thật), khối 2 nhánh loại trừ lẫn nhau ("the_tich"/
"be_mat"), công thức R=Q1/30 bổ sung cho id=22 (bị thiếu trong .docx do vốn là
equation object không convert được sang text), cảnh báo KHÔNG tự bịa số liệu
từ biểu đồ Hình C.1-C.4 (id=29 — điểm dễ sai nhất), và xác nhận KHÔNG áp dụng
KHONG_UOC_LUONG_KHOANG_CACH/TOA_DO_TRUC_KHOANG_CACH (khác B15 — owner quyết
định B16 không có tiêu chí khoảng cách giữa 2 thiết bị nào).

Không gọi AI thật — chỉ kiểm tra chuỗi system prompt sinh ra."""

from app.services import mdc_filler
from app.services.bot_chua_chay_reader import SYSTEM_PROMPT, _BE_MAT_IDS, _EXPECTED_IDS, _THE_TICH_IDS

_THE_TICH_LIST = [19, 20, 21, 22, 23, 24]
_BE_MAT_LIST = [27, 28, 29]
_COMMON_LIST = [1, 2, 6, 7, 9, 10, 12, 13, 14, 15, 31, 32, 34, 35, 37, 38, 39, 40, 42, 43, 44, 46, 47, 48]
_ALL_EXPECTED_LIST = sorted(_COMMON_LIST + _THE_TICH_LIST + _BE_MAT_LIST)


def test_expected_ids_match_real_docx():
    rows = mdc_filler.load_criteria_rows("bot_chua_chay")
    assert [r["id"] for r in rows] == _ALL_EXPECTED_LIST
    assert len(_ALL_EXPECTED_LIST) == 33
    assert _EXPECTED_IDS == set(_ALL_EXPECTED_LIST)
    assert tuple(_THE_TICH_IDS) == tuple(_THE_TICH_LIST)
    assert tuple(_BE_MAT_IDS) == tuple(_BE_MAT_LIST)


def test_template_has_exactly_1_table_after_extra_removed():
    """File goc convert tu .doc co 1 bang thua 1x2 dung TRUOC bang tieu chi
    that - phai da bi xoa, doc.tables[0] phai la bang that 49 dong (48 tieu
    chi + header)."""
    from docx import Document
    doc = Document(mdc_filler.TEMPLATE_PATHS["bot_chua_chay"])
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 49
    header = [c.text.strip() for c in doc.tables[0].rows[0].cells]
    assert header[0] == "TT"
    assert "Kết luận" in header[-1]


def test_criteria_list_contains_all_33_ids():
    for i in _ALL_EXPECTED_LIST:
        assert f"id={i} |" in SYSTEM_PROMPT


def test_nhanh_block_present_with_exact_wording():
    assert '"x - Không áp dụng"' in SYSTEM_PROMPT
    assert "THEO THỂ TÍCH" in SYSTEM_PROMPT
    assert "BỀ MẶT" in SYSTEM_PROMPT
    assert '"nhanh": "the_tich" | "be_mat"' in SYSTEM_PROMPT


def test_toc_do_xa_r_formula_supplied():
    """id=22 thieu cong thuc R=Q1/30 trong .docx (equation object) - phai duoc
    bo sung rieng trong prompt."""
    assert "R = Q1/30" in SYSTEM_PROMPT


def test_k_coefficients_not_confused():
    assert "K1.V + K2.AS + K3.AL + K4.Rv.t" in SYSTEM_PROMPT or "K1-K4" in SYSTEM_PROMPT
    assert "K5 = 1,2 kg/m³" in SYSTEM_PROMPT
    assert "KHÔNG áp K5 vào tính Q1" in SYSTEM_PROMPT


def test_hinh_c_nomogram_caveat_present():
    """Diem de sai nhat cua B16 - AI khong duoc tu doan so lieu tu bieu do."""
    assert "Hình C.1" in SYSTEM_PROMPT and "Hình C.4" in SYSTEM_PROMPT
    assert "KHÔNG THỂ trích xuất chính xác bằng đọc text/OCR" in SYSTEM_PROMPT
    assert "TUYỆT ĐỐI KHÔNG được tự đoán/bịa" in SYSTEM_PROMPT
    assert "cần tra biểu đồ thủ công" in SYSTEM_PROMPT


def test_no_distance_rules_applied_unlike_b15():
    """Owner quyet dinh B16 khong co tieu chi 'khoang cach giua 2 thiet bi'
    nao - KHONG duoc co KHONG_UOC_LUONG_KHOANG_CACH/TOA_DO_TRUC_KHOANG_CACH."""
    assert "tại vị trí trục" not in SYSTEM_PROMPT
    assert "KHÔNG tự ước lượng khoảng cách" not in SYSTEM_PROMPT


def test_powder_vs_foam_distinction_present():
    assert "KHÁC HẲN hệ thống chữa cháy bằng BỌT" in SYSTEM_PROMPT


def test_shared_checklist_reused_not_duplicated():
    from app.services.ai_reader_common import NHOM_II_MAU_THUAN_CHECKLIST
    assert SYSTEM_PROMPT.count(NHOM_II_MAU_THUAN_CHECKLIST.strip()[:40]) == 1


def test_no_khibotsolkhi_style_he_thong_classification():
    assert '"he_thong"' not in SYSTEM_PROMPT
    assert 'NẾU he_thong' not in SYSTEM_PROMPT
