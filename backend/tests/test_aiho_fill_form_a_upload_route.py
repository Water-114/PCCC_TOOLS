"""Batch 5A Pha 3 Buoc 5 — test route POST /api/aiho/fill-form-a-upload: nhan
multipart form (file .docx Form A trong + session_id + hang_muc_json), goi AI
text-only (mock hoan toan), tra ve .docx da dien base64. Kiem tra ca gioi han
file/form cua phien (+1 file, +1 form giong het cac route AI thuc su khac,
KHONG phai route "chi doc lai" nhu /export-*)."""

import base64
import io
import json

from docx import Document

from app.models import AIHO_API_NAME, HoSoSession, count_usage_today
from app.providers.base import GenerationResult
from app.services import credits
from app.services.mdc_filler import COL_DOI_CHIEU, COL_KET_LUAN, COL_KHOAN_DIEU, COL_QUY_DINH, COL_THIET_KE


def _register_login_and_grant(client, email="formaup1@pccc.local", password="matkhau123", amount=5):
    client.post("/api/auth/register", json={"email": email, "password": password})
    resp = client.post("/api/auth/login", json={"email": email, "password": password})
    data = resp.get_json()
    token, user_id = data["token"], data["user"]["id"]
    if amount:
        credits.grant_credits(user_id, amount, credits.CREDIT_REASON_EMAIL_VERIFICATION, note="test setup")
    return token, user_id


def _open_session(client, token):
    resp = client.post("/api/aiho/session/open", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 200, resp.get_json()
    return resp.get_json()["session_id"]


def _fake_form_a_bytes():
    doc = Document()
    table = doc.add_table(rows=2, cols=6)
    table.rows[0].cells[0].text = "TT"
    table.rows[0].cells[COL_DOI_CHIEU].text = "Đối tượng"
    table.rows[0].cells[COL_THIET_KE].text = "Nội dung thiết kế"
    table.rows[0].cells[COL_QUY_DINH].text = "Quy định"
    table.rows[0].cells[COL_KHOAN_DIEU].text = "Khoản/Điều"
    table.rows[0].cells[COL_KET_LUAN].text = "Kết luận"
    table.rows[1].cells[COL_DOI_CHIEU].text = "Hệ thống 1"
    table.rows[1].cells[COL_QUY_DINH].text = "Phải đáp ứng quy định 1"
    table.rows[1].cells[COL_KHOAN_DIEU].text = "Điều 1"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class FakeProvider:
    def __init__(self, name="fake", model="fake-model", fn=None, exc=None):
        self.name = name
        self.model = model
        self.fn = fn
        self.exc = exc

    def generate(self, prompt):
        if self.exc is not None:
            raise self.exc
        return self.fn(prompt)


def _valid_payload_fn(prompt):
    return json.dumps({"items": [{"id": 1, "noi_dung_thiet_ke": "Đã đối chiếu đủ căn cứ", "ket_luan": "dat"}]})


def _upload(client, token, session_id, extra_form=None):
    from unittest.mock import patch
    provider = FakeProvider(fn=_valid_payload_fn)
    data = {
        "session_id": str(session_id),
        "hang_muc_json": json.dumps([{"ten_he_thong": "Báo cháy", "items": [{"noi_dung_thiet_ke": "x", "ket_luan": "dat"}]}]),
        "file": (io.BytesIO(_fake_form_a_bytes()), "FormA_trong.docx"),
    }
    if extra_form:
        data.update(extra_form)
    with patch("app.routes.aiho.get_provider", return_value=provider):
        return client.post(
            "/api/aiho/fill-form-a-upload",
            data=data,
            headers={"Authorization": f"Bearer {token}"},
            content_type="multipart/form-data",
        )


def test_requires_login(client):
    resp = client.post("/api/aiho/fill-form-a-upload", data={"session_id": "1"}, content_type="multipart/form-data")
    assert resp.status_code == 401


def test_missing_session_id_returns_400(client):
    token, _ = _register_login_and_grant(client, email="formaup2@pccc.local")
    resp = client.post(
        "/api/aiho/fill-form-a-upload",
        data={"file": (io.BytesIO(_fake_form_a_bytes()), "x.docx")},
        headers={"Authorization": f"Bearer {token}"},
        content_type="multipart/form-data",
    )
    assert resp.status_code == 400


def test_missing_file_returns_400(client):
    token, _ = _register_login_and_grant(client, email="formaup3@pccc.local")
    session_id = _open_session(client, token)
    resp = client.post(
        "/api/aiho/fill-form-a-upload",
        data={"session_id": str(session_id)},
        headers={"Authorization": f"Bearer {token}"},
        content_type="multipart/form-data",
    )
    assert resp.status_code == 400


def test_non_docx_file_returns_400(client):
    token, _ = _register_login_and_grant(client, email="formaup4@pccc.local")
    session_id = _open_session(client, token)
    resp = client.post(
        "/api/aiho/fill-form-a-upload",
        data={"session_id": str(session_id), "file": (io.BytesIO(b"not a docx file"), "x.docx")},
        headers={"Authorization": f"Bearer {token}"},
        content_type="multipart/form-data",
    )
    assert resp.status_code == 400


def test_invalid_hang_muc_json_returns_400(client):
    token, _ = _register_login_and_grant(client, email="formaup5@pccc.local")
    session_id = _open_session(client, token)
    resp = client.post(
        "/api/aiho/fill-form-a-upload",
        data={
            "session_id": str(session_id),
            "file": (io.BytesIO(_fake_form_a_bytes()), "x.docx"),
            "hang_muc_json": "{not valid json",
        },
        headers={"Authorization": f"Bearer {token}"},
        content_type="multipart/form-data",
    )
    assert resp.status_code == 400


def test_valid_payload_returns_200_and_fills_docx(client):
    token, _ = _register_login_and_grant(client, email="formaup6@pccc.local")
    session_id = _open_session(client, token)
    resp = _upload(client, token, session_id)
    assert resp.status_code == 200, resp.get_json()
    data = resp.get_json()
    assert data["filename"].endswith(".docx")
    docx_bytes = base64.b64decode(data["base64"])
    doc = Document(io.BytesIO(docx_bytes))
    assert doc.tables[0].rows[1].cells[COL_KET_LUAN].text == "Đạt"
    assert data["ho_so"]["files_used"] == 1
    assert data["ho_so"]["forms_used"] == 1


def test_valid_payload_writes_success_usage_log(client):
    token, user_id = _register_login_and_grant(client, email="formaup7@pccc.local")
    session_id = _open_session(client, token)
    before = count_usage_today(user_id, AIHO_API_NAME)
    resp = _upload(client, token, session_id)
    assert resp.status_code == 200
    after = count_usage_today(user_id, AIHO_API_NAME)
    assert after == before + 1


def test_other_users_session_returns_404(client):
    token1, _ = _register_login_and_grant(client, email="formaup8a@pccc.local")
    session_id = _open_session(client, token1)
    token2, _ = _register_login_and_grant(client, email="formaup8b@pccc.local")
    resp = _upload(client, token2, session_id)
    assert resp.status_code == 404


def test_ai_error_returns_502_and_still_counts_files_used(client):
    from unittest.mock import patch
    token, _ = _register_login_and_grant(client, email="formaup9@pccc.local")
    session_id = _open_session(client, token)
    provider = FakeProvider(fn=lambda prompt: "not valid json at all")
    data = {
        "session_id": str(session_id),
        "hang_muc_json": "[]",
        "file": (io.BytesIO(_fake_form_a_bytes()), "FormA_trong.docx"),
    }
    with patch("app.routes.aiho.get_provider", return_value=provider):
        resp = client.post(
            "/api/aiho/fill-form-a-upload",
            data=data,
            headers={"Authorization": f"Bearer {token}"},
            content_type="multipart/form-data",
        )
    assert resp.status_code == 502
    session = HoSoSession.query.get(session_id)
    assert session.files_used == 1
