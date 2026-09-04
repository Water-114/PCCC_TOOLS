"""Batch 5A Pha 3 Buoc 3 - xuat Cong van huong dan (.docx that, dung file
mau owner cung cap). Kiem tra: (a) tieu de nhom CO kien nghi con lai dung so
dong; (b) tieu de nhom KHONG co kien nghi bi xoa han; (c) MERGEFIELD thay
dung gia tri; (d) field quy_mo de trong -> "[yêu cầu nhập thông tin]" mau do."""

from docx import Document

from app.services.cong_van_huong_dan_docx import (
    CongVanHuongDanError,
    build_cong_van_huong_dan_docx,
)
from app.services.docx_mergefield import YEU_CAU_NHAP_THONG_TIN
from app.services.mdc_filler import MAU_DO

EMPTY_KIEN_NGHI = {"I_chua_the_hien": [], "II_chua_thong_nhat": [], "III_chua_phu_hop": [], "IV_de_xuat_bo_sung": []}


def _kien_nghi(**overrides):
    kn = dict(EMPTY_KIEN_NGHI)
    kn.update(overrides)
    return kn


def _open_result(docx_bytes):
    import io
    return Document(io.BytesIO(docx_bytes))


def test_build_cong_van_huong_dan_docx_full_scenario():
    session_data = {
        "quy_mo": {
            "tenCongTrinh": "Nhà máy Test ABC",
            "chuDauTu": "Công ty XYZ",
            "diaDiemXayDung": "Số 1 đường Test",
            "diaChiChuDauTu": None,  # de trong -> yeu cau nhap thong tin, mau do
            "donViTuVanThietKe": "Công ty tư vấn Test",
            "soNgayPC11": "số 99 ngày 01/01/2026",
            "maHoSo": "HS-TEST-01",
        }
    }
    hang_muc_list = [
        {"slot": "quy_mo", "kien_nghi": _kien_nghi(
            I_chua_the_hien=["Bổ sung số liệu diện tích A."],
            IV_de_xuat_bo_sung=["Đề xuất bổ sung khối tích B."],
        )},
        {"slot": "baochay", "kien_nghi": _kien_nghi(
            I_chua_the_hien=["Câu báo cháy 1.", "Câu báo cháy 2."],
            III_chua_phu_hop=["Câu báo cháy 3."],
        )},
        {"slot": "ccnuoc", "kien_nghi": _kien_nghi(I_chua_the_hien=["Câu chữa cháy nước 1."])},
        {"slot": "dienpccc", "kien_nghi": _kien_nghi()},  # rong -> nhom "dien" bi xoa het
        # "densucco" KHONG dinh kem -> nhom "khac" cung phai bi xoa het
    ]

    docx_bytes = build_cong_van_huong_dan_docx(session_data, hang_muc_list)
    doc = _open_result(docx_bytes)
    all_texts = [p.text.strip() for p in doc.paragraphs]

    # (a) tieu de CO kien nghi con lai, dung so dong "- ..."
    assert "Thông tin công trình" in all_texts
    assert "- Bổ sung số liệu diện tích A." in all_texts
    assert "- Đề xuất bổ sung khối tích B." in all_texts

    assert "Hệ thống báo cháy:" in all_texts
    assert "- Câu báo cháy 1." in all_texts
    assert "- Câu báo cháy 2." in all_texts
    assert "- Câu báo cháy 3." in all_texts

    assert "Hệ thống chữa cháy:" in all_texts
    assert "- Câu chữa cháy nước 1." in all_texts

    # (b) tieu de KHONG co kien nghi da bi xoa han (khong con trong doc.paragraphs)
    assert "Hệ thống điện:" not in all_texts
    assert "Các hệ thống, phương tiện PCCC khác:" not in all_texts

    # (c) MERGEFIELD da thay dung gia tri test dua vao (kiem tra ca body + table letterhead)
    full_text = "\n".join(all_texts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text += "\n" + p.text
    assert "Nhà máy Test ABC" in full_text
    assert "Công ty XYZ" in full_text
    assert "Số 1 đường Test" in full_text
    assert "Công ty tư vấn Test" in full_text
    assert "số 99 ngày 01/01/2026" in full_text
    assert "HS-TEST-01" in full_text

    # (d) field de trong -> "[yêu cầu nhập thông tin]" mau do EE0000
    found_placeholder = False
    for p in doc.paragraphs:
        if YEU_CAU_NHAP_THONG_TIN in p.text:
            for r in p.runs:
                if r.text.strip() == YEU_CAU_NHAP_THONG_TIN:
                    found_placeholder = True
                    assert r.font.color.rgb == MAU_DO
    assert found_placeholder, "Khong tim thay placeholder [yêu cầu nhập thông tin] cho dia_chi_chu_dau_tu de trong"


def test_build_cong_van_huong_dan_docx_missing_template_raises():
    import app.services.cong_van_huong_dan_docx as mod
    original = mod.TEMPLATE_PATH
    try:
        mod.TEMPLATE_PATH = mod.mdc_filler.TEMPLATES_DIR / "khong_ton_tai_xyz.docx"
        try:
            build_cong_van_huong_dan_docx({"quy_mo": {}}, [])
            assert False, "Phai raise CongVanHuongDanError khi thieu file mau"
        except CongVanHuongDanError:
            pass
    finally:
        mod.TEMPLATE_PATH = original
