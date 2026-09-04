"""Batch 5A Pha 3 Buoc 4 - xuat Bao cao tham dinh PCCC (11 muc, .docx that,
dung file mau owner cung cap). Kiem tra: (a) muc 1-2-3 dien dung du KHONG co
hang_muc nao lien quan (chi tu quy_mo) - dac biet muc 1/4 co "mo ta hien
trang" nam TREN CAC DOAN RIENG trong file mau that (khac cau truc gia dinh
ban dau), phai gop dung ve 1 doan; (b) muc 5 co bullet (tach tu densucco
gop), muc 6/9 co cau ghi chu "xem muc 5"; (c) muc 8 KHONG co slot lien quan
-> "khong.", KHONG co khoi "Kien nghi:"; (d) muc 8 CO du lieu (gia lap
khibot) du template mau dang la "khong." -> code TU CHEN duoc khoi "Kien
nghi:" + bullet moi (phep thu quan trong nhat, xac nhan khong phu thuoc cau
truc cu the cua vi du Le Quynh)."""

import io

from docx import Document

from app.services.bao_cao_tham_dinh_docx import (
    BaoCaoThamDinhError,
    build_bao_cao_tham_dinh_docx,
)

EMPTY_KIEN_NGHI = {"I_chua_the_hien": [], "II_chua_thong_nhat": [], "III_chua_phu_hop": [], "IV_de_xuat_bo_sung": []}


def _kn(**overrides):
    d = dict(EMPTY_KIEN_NGHI)
    d.update(overrides)
    return d


def _session_data():
    return {
        "quy_mo": {
            "occ": "nhahang",
            "tenCongTrinh": "Nhà hàng Test ABC",
            "chuDauTu": "Công ty Test",
            "diaDiemXayDung": "123 Đường Test",
            "diaChiChuDauTu": "456 Đường Chủ",
            "donViTuVanThietKe": "Công ty tư vấn Test",
            "maHoSo": "HS-TEST-01",
            "tongMucDauTu": "5.000.000.000 đồng",
            "hazard": "C",
            "bacChiuLua": "II",
            "capNguyHiemChayKetCau": "S0",
            "floors": 4,
            "basements": 1,
            "totalArea": 3000,
            "hFire": 13.5,
            "thanhPhanHoSo": ["Văn bản PC11 số 01.", "Giấy chứng nhận quyền sử dụng đất số X."],
        }
    }


def _hang_muc_base():
    return [
        {"slot": "baochay", "tong_ket": "Hệ thống báo cháy địa chỉ, 5 zone.", "kien_nghi": _kn(I_chua_the_hien=["Câu báo cháy 1."])},
        {"slot": "ccnuoc", "tong_ket": "Hệ thống họng nước + sprinkler đầy đủ.", "kien_nghi": _kn(III_chua_phu_hop=["Câu nước 1."])},
        {"slot": "dienpccc", "tong_ket": "Điện PCCC đạt yêu cầu cơ bản.", "kien_nghi": _kn(I_chua_the_hien=["Câu điện 1.", "Câu điện 2."])},
        {
            "slot": "densucco",
            "tong_ket": (
                "Đèn chiếu sáng sự cố, đèn chỉ dẫn thoát nạn: bố trí đầy đủ trên đường thoát nạn. "
                "Bình chữa cháy xách tay/xe đẩy: bố trí tại các tầng."
            ),
            "kien_nghi": _kn(
                I_chua_the_hien=["Bổ sung sơ đồ chỉ dẫn thoát nạn tầng 3."],
                III_chua_phu_hop=["Bổ sung bình chữa cháy tầng 2."],
            ),
        },
        {"slot": "quy_mo", "kien_nghi": _kn(II_chua_thong_nhat=["Số liệu diện tích chưa thống nhất."])},
    ]


def _find(texts, prefix):
    return next(t for t in texts if t.startswith(prefix))


def _idx(texts, prefix):
    return next(i for i, t in enumerate(texts) if t.startswith(prefix))


def test_muc_1_2_3_dien_dung_tu_quy_mo_khong_can_hang_muc():
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), [])
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    muc1 = _find(texts, "1. Tính pháp lý của hồ sơ:")
    assert "Nhà hàng Test ABC" in muc1
    # xac nhan KHONG con doan mo ta rieng le nao sot lai (file mau that co 2
    # doan rieng sau tieu de muc 1 - phai duoc gop het vao 1 doan nay)
    idx1 = _idx(texts, "1. Tính pháp lý của hồ sơ:")
    assert texts[idx1 + 1].startswith("2. Thành phần và số lượng hồ sơ:")

    muc2 = _find(texts, "2. Thành phần và số lượng hồ sơ:")
    assert muc2 == "2. Thành phần và số lượng hồ sơ: phù hợp theo quy định tại khoản 4 Điều 9 Nghị định số 105/2025/NĐ-CP, gồm:"
    idx2 = _idx(texts, "2. Thành phần và số lượng hồ sơ:")
    assert texts[idx2 + 1] == "- Văn bản PC11 số 01."
    assert texts[idx2 + 2] == "- Giấy chứng nhận quyền sử dụng đất số X."

    muc3 = _find(texts, "3. Hạng nguy hiểm cháy")
    assert "hạng nguy hiểm cháy nổ C" in muc3
    assert "bậc chịu lửa II" in muc3
    assert "cấp nguy hiểm cháy kết cấu S0" in muc3


def test_muc_4_gop_dung_mo_ta_rieng_va_giu_kien_nghi_quy_mo():
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), _hang_muc_base())
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    idx4 = _idx(texts, "4. Quy mô")
    assert "Số tầng:" in texts[idx4]
    assert "Tổng diện tích sàn: 3.000 m²" in texts[idx4]
    assert texts[idx4 + 1] == "Kiến nghị:"
    assert texts[idx4 + 2] == "- Số liệu diện tích chưa thống nhất."


def test_muc_5_6_9_densucco_tach_dung_va_ghi_chu_xem_muc_5():
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), _hang_muc_base())
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    idx5 = _idx(texts, "5. Hệ thống đèn chiếu sáng")
    assert "bố trí đầy đủ trên đường thoát nạn" in texts[idx5]
    assert texts[idx5 + 1] == "Kiến nghị:"
    assert texts[idx5 + 2] == "- Bổ sung sơ đồ chỉ dẫn thoát nạn tầng 3."
    assert texts[idx5 + 3] == "- Bổ sung bình chữa cháy tầng 2."

    muc6 = _find(texts, "6. Hệ thống báo cháy")
    assert "xem đối chiếu gộp tại mục 5" in muc6

    muc9 = _find(texts, "9. Trang bị phương tiện chữa cháy khác:")
    assert "bố trí tại các tầng" in muc9
    assert "xem mục 5" in muc9
    # muc 9 KHONG co bullet kien nghi rieng (khong tach duoc tu densucco gop)
    idx9 = _idx(texts, "9. Trang bị phương tiện chữa cháy khác:")
    assert texts[idx9 + 1].startswith("10. Hệ thống điện")


def test_muc_8_rong_khi_khong_co_slot_lien_quan():
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), _hang_muc_base())
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    assert "8. Hệ thống chữa cháy bằng bọt, bột, khí: không." in texts
    idx8 = _idx(texts, "8. Hệ thống chữa cháy bằng bọt, bột, khí:")
    # doan ngay sau la tieu de muc 9, KHONG phai "Kien nghi:"
    assert texts[idx8 + 1].startswith("9. Trang bị phương tiện chữa cháy khác:")


def test_muc_8_tu_chen_khoi_kien_nghi_moi_khi_co_du_lieu_that():
    """Phep thu quan trong nhat: template mau Le Quynh dang la '8. ...: khong.'
    (khong co khoi 'Kien nghi:' nao san) nhung du lieu THAT cua du an nay co
    kien nghi cho muc 8 (gia lap qua slot 'khibot') - code phai TU CHEN duoc
    1 khoi 'Kien nghi:' + bullet MOI, khong phu thuoc cau truc co san."""
    hang_muc = _hang_muc_base() + [{
        "slot": "khibot",
        "tong_ket": "Hệ thống khí hoá lỏng FM200 bố trí đầy đủ cho phòng server.",
        "kien_nghi": _kn(I_chua_the_hien=["Bổ sung tính toán nồng độ thiết kế khí FM200."]),
    }]
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), hang_muc)
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    idx8 = _idx(texts, "8. Hệ thống chữa cháy bằng bọt, bột, khí:")
    assert "FM200" in texts[idx8]
    assert "không." not in texts[idx8]
    assert texts[idx8 + 1] == "Kiến nghị:"
    assert texts[idx8 + 2] == "- Bổ sung tính toán nồng độ thiết kế khí FM200."
    # muc 9 van con nguyen ngay sau, khong bi anh huong
    assert texts[idx8 + 3].startswith("9. Trang bị phương tiện chữa cháy khác:")


def test_mergefields_thay_dung_gia_tri():
    docx_bytes = build_bao_cao_tham_dinh_docx(_session_data(), _hang_muc_base())
    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + "\n".join(p.text for p in cell.paragraphs)

    assert "Nhà hàng Test ABC" in full_text
    assert "Công ty Test" in full_text
    assert "123 Đường Test" in full_text
    assert "456 Đường Chủ" in full_text
    assert "Công ty tư vấn Test" in full_text
    assert "HS-TEST-01" in full_text
    assert "5.000.000.000 đồng" in full_text


def test_missing_template_raises():
    import app.services.bao_cao_tham_dinh_docx as mod
    original = mod.TEMPLATE_PATH
    try:
        mod.TEMPLATE_PATH = mod.mdc_filler.TEMPLATES_DIR / "khong_ton_tai_xyz.docx"
        try:
            build_bao_cao_tham_dinh_docx({"quy_mo": {}}, [])
            assert False, "Phải raise BaoCaoThamDinhError khi thiếu file mẫu"
        except BaoCaoThamDinhError:
            pass
    finally:
        mod.TEMPLATE_PATH = original
