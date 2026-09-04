"""Batch 5A Pha 3 Buoc 5 - dien Form A do nguoi dung TU DINH KEM (blank
template) dua tren findings da co trong phien (KHONG doc lai ban ve). Lan
DAU TIEN app goi AI dang text-only (khong kem anh/PDF) - dung fake .docx tu
dung bang python-docx (khong can file mau that), mock provider.generate()."""

import io
import json

from docx import Document

from app.services.form_a_upload import FormAUploadError, dien_form_a_upload
from app.services.mdc_filler import (
    COL_DOI_CHIEU,
    COL_KET_LUAN,
    COL_KHOAN_DIEU,
    COL_QUY_DINH,
    COL_THIET_KE,
    MAU_DO,
)


class FakeProvider:
    name = "fake"
    model = "fake-model"

    def __init__(self, fn):
        self.fn = fn
        self.calls = []

    def generate(self, prompt):
        self.calls.append(prompt)
        return self.fn(prompt)


def _build_fake_form_a_docx(n_rows=3):
    """Dung 1 file .docx gia lap dung bo cuc 6 cot cua Form A/MDC that
    (TT | doi_chieu | thiet_ke | quy_dinh | khoan_dieu | ket_luan)."""
    doc = Document()
    table = doc.add_table(rows=n_rows + 1, cols=6)
    table.rows[0].cells[0].text = "TT"
    table.rows[0].cells[COL_DOI_CHIEU].text = "Đối tượng"
    table.rows[0].cells[COL_THIET_KE].text = "Nội dung thiết kế"
    table.rows[0].cells[COL_QUY_DINH].text = "Quy định"
    table.rows[0].cells[COL_KHOAN_DIEU].text = "Khoản/Điều"
    table.rows[0].cells[COL_KET_LUAN].text = "Kết luận"
    for i in range(1, n_rows + 1):
        table.rows[i].cells[COL_DOI_CHIEU].text = f"Hệ thống {i}"
        table.rows[i].cells[COL_QUY_DINH].text = f"Phải đáp ứng quy định {i}"
        table.rows[i].cells[COL_KHOAN_DIEU].text = f"Điều {i}"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_dien_form_a_upload_dien_dung_cot_3_5_va_to_do():
    file_bytes = _build_fake_form_a_docx(n_rows=3)

    def payload_fn(prompt):
        assert "id=1" in prompt and "id=2" in prompt and "id=3" in prompt
        return json.dumps({"items": [
            {"id": 1, "noi_dung_thiet_ke": "Đã có báo cháy địa chỉ 5 zone", "ket_luan": "dat"},
            # LUU Y: khong dung cum tu "khong su dung/khong co.../khong trang bi..."
            # o day - ket_luan_linter.py (Pha 0) se TU DONG doi chua_dat/chua_the_hien
            # -> khong_ap_dung khi gap cac cum tu do (dung nhu thiet ke), lam sai
            # muc dich test nay (dang can kiem "chua_dat" hien thi "KN" binh thuong).
            {"id": 2, "noi_dung_thiet_ke": "Sprinkler bố trí chưa đủ khoảng cách theo quy định", "ket_luan": "chua_dat"},
            {"id": 3, "noi_dung_thiet_ke": "Chưa đối chiếu được từ dữ liệu đã đọc trong bộ hồ sơ này — cần rà thêm", "ket_luan": "chua_the_hien"},
        ]})

    provider = FakeProvider(payload_fn)
    hang_muc_digest = [{"ten_he_thong": "Báo cháy tự động", "items": [{"noi_dung_thiet_ke": "x", "ket_luan": "dat"}]}]
    result_bytes = dien_form_a_upload(file_bytes, hang_muc_digest, {"occ": "chungcu"}, provider)

    doc = Document(io.BytesIO(result_bytes))
    table = doc.tables[0]

    row1_thiet_ke = table.rows[1].cells[COL_THIET_KE]
    assert row1_thiet_ke.text == "Đã có báo cháy địa chỉ 5 zone"
    assert row1_thiet_ke.paragraphs[0].runs[0].font.color.rgb == MAU_DO
    row1_ket_luan = table.rows[1].cells[COL_KET_LUAN]
    assert row1_ket_luan.text == "Đạt"
    # "dat" khong to do o cot Ket luan
    assert not row1_ket_luan.paragraphs[0].runs or row1_ket_luan.paragraphs[0].runs[0].font.color.rgb != MAU_DO

    row2_ket_luan = table.rows[2].cells[COL_KET_LUAN]
    assert row2_ket_luan.text == "KN"
    assert row2_ket_luan.paragraphs[0].runs[0].font.color.rgb == MAU_DO

    row3_ket_luan = table.rows[3].cells[COL_KET_LUAN]
    assert row3_ket_luan.text == "KN"
    assert row3_ket_luan.paragraphs[0].runs[0].font.color.rgb == MAU_DO

    # cot doi_chieu/quy_dinh/khoan_dieu giu nguyen 100%, khong bi doi
    assert table.rows[1].cells[COL_DOI_CHIEU].text == "Hệ thống 1"
    assert table.rows[1].cells[COL_QUY_DINH].text == "Phải đáp ứng quy định 1"
    assert table.rows[1].cells[COL_KHOAN_DIEU].text == "Điều 1"


def test_dien_form_a_upload_hang_muc_digest_rong_van_hop_le():
    """Chua doc hang muc nao trong phien - AI van nhan duoc digest rong hop
    le, khong loi, van dien duoc (theo huong "chua doi chieu duoc")."""
    file_bytes = _build_fake_form_a_docx(n_rows=2)

    def payload_fn(prompt):
        assert "Chưa có hạng mục nào đã đọc trong phiên này" in prompt
        return json.dumps({"items": [
            {"id": 1, "noi_dung_thiet_ke": "Chưa đối chiếu được từ dữ liệu đã đọc trong bộ hồ sơ này — cần rà thêm", "ket_luan": "chua_the_hien"},
            {"id": 2, "noi_dung_thiet_ke": "Chưa đối chiếu được từ dữ liệu đã đọc trong bộ hồ sơ này — cần rà thêm", "ket_luan": "chua_the_hien"},
        ]})

    provider = FakeProvider(payload_fn)
    result_bytes = dien_form_a_upload(file_bytes, [], {}, provider)
    doc = Document(io.BytesIO(result_bytes))
    table = doc.tables[0]
    assert table.rows[1].cells[COL_KET_LUAN].text == "KN"
    assert table.rows[2].cells[COL_KET_LUAN].text == "KN"


def test_dien_form_a_upload_khong_co_bang_raise_error():
    doc = Document()
    doc.add_paragraph("Khong co bang nao trong file nay")
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    provider = FakeProvider(lambda prompt: json.dumps({"items": []}))
    try:
        dien_form_a_upload(file_bytes, [], {}, provider)
        assert False, "Phải raise FormAUploadError khi file không có bảng"
    except FormAUploadError as exc:
        assert "bảng" in str(exc)


def test_dien_form_a_upload_file_hong_raise_error():
    provider = FakeProvider(lambda prompt: json.dumps({"items": []}))
    try:
        dien_form_a_upload(b"khong phai file docx that", [], {}, provider)
        assert False, "Phải raise FormAUploadError khi file hỏng"
    except FormAUploadError:
        pass


def test_dien_form_a_upload_khong_goi_lai_ai_khi_khong_co_dong_tieu_chi():
    """Bang chi co dong tieu de, khong co dong tieu chi nao - phai raise
    loi ro rang, KHONG goi AI (lang phi luot goi)."""
    doc = Document()
    table = doc.add_table(rows=1, cols=6)
    table.rows[0].cells[0].text = "TT"
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    provider = FakeProvider(lambda prompt: json.dumps({"items": []}))
    try:
        dien_form_a_upload(file_bytes, [], {}, provider)
        assert False, "Phải raise FormAUploadError khi không có dòng tiêu chí nào"
    except FormAUploadError:
        pass
    assert provider.calls == []
