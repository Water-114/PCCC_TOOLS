"""Dung file mau .docx that de xuat "Bao cao tham dinh PCCC" (11 muc trong
Phan III) - doi chieu tai lieu "Bo quy tac doc ban ve va dien MDC" 03/9/2026.
Dung lai docx_mergefield.replace_mergefield() (Buoc 3) - KHONG viet lai.

QUAN TRONG - da doc truc tiep file mau that (khong doan): 11 muc KHONG dong
nhat cau truc nhu gia dinh ban dau - muc 1 va muc 4 co doan "mo ta hien
trang" nam TREN (CAC) DOAN RIENG NGAY SAU doan tieu de (khac muc 3/5/6/7/8/
9/10/11 co mo ta nam CHUNG 1 doan voi tieu de). Vi vay _thay_heading_va_kien_nghi()
xu ly DONG: gom moi doan "mo ta" phu (neu co) ngay sau tieu de truoc khi ghi
de, khong gia dinh truoc 1 doan duy nhat."""

import copy
import io

from docx import Document
from docx.oxml.ns import qn

from . import mdc_filler
from .densucco_reader import FORMS as _DENSUCCO_FORMS
from .docx_mergefield import replace_mergefield
from .quy_mo_store import _fmt as _fmt_so  # tai dung dinh dang so (dau . ngan nghin) - khong viet trung
from .tham_dinh import ThamDinhInputError, evaluate_tham_dinh

TEMPLATE_PATH = mdc_filler.TEMPLATES_DIR / "bao_cao_tham_dinh.docx"
FILENAME = "Bao_cao_tham_dinh_PCCC.docx"

# Nguyen van phan dau (giu NGUYEN) cua 11 muc - copy CHINH XAC tu file mau
# that (da doc truc tiep, khong go lai theo tri nho) de do dung paragraph.
_MUC_PREFIX = {
    1: "1. Tính pháp lý của hồ sơ:",
    2: "2. Thành phần và số lượng hồ sơ: phù hợp theo quy định tại khoản 4 Điều 9 Nghị định số 105/2025/NĐ-CP, gồm:",
    3: "3. Hạng nguy hiểm cháy, nổ, dự kiến bậc chịu lửa, cấp nguy hiểm cháy kết cấu của nhà, nhóm nguy hiểm cháy của nhà:",
    4: "4. Quy mô (diện tích, khối tích, số tầng, chiều cao PCCC, công suất…):",
    5: "5. Hệ thống đèn chiếu sáng sự cố và chỉ dẫn thoát nạn:",
    6: "6. Hệ thống báo cháy tự động:",
    7: "7. Hệ thống chữa cháy bằng nước gồm:",
    8: "8. Hệ thống chữa cháy bằng bọt, bột, khí:",
    9: "9. Trang bị phương tiện chữa cháy khác:",
    10: "10. Hệ thống điện cấp cho phòng cháy và chữa cháy:",
    11: "11. Các hệ thống khác có liên quan:",
}

# Muc nao dung cau truc "heading -> bullet THANG" (nhu VBHD, KHONG co doan
# "Kien nghi:" xen giua) - da xac nhan qua file that CHI muc 2. Cac muc con
# lai dung cau truc "heading (+ mo ta) -> [Kien nghi: -> bullet]" (co the co
# hoac khong co khoi Kien nghi tuy du lieu THAT, khong gia dinh truoc).
_MUC_BULLET_TRUC_TIEP = {2}

_KIEN_NGHI_LABEL = "Kiến nghị:"
_PHAN_IV_HEADING = "IV. NHẬN XÉT, ĐỀ XUẤT"

# Dung de biet DIEM DUNG khi gom cac doan "mo ta" phu ngay sau 1 tieu de muc
# (xem _thay_heading_va_kien_nghi) - la tap CAC CHUOI DA BIET TRUOC (11 tieu
# de + nhan Kien nghi + tieu de Phan IV), KHONG doan bang regex long leo.
_STOP_EXACT = {_KIEN_NGHI_LABEL}
_STOP_PREFIXES = tuple(_MUC_PREFIX.values()) + (_PHAN_IV_HEADING,)


def _la_diem_dung(text: str) -> bool:
    """True neu doan nay la 1 "moc" da biet truoc (nhan Kien nghi, hoac tieu
    de cua 1 trong 11 muc/Phan IV) - dung startswith (KHONG phai ==) vi 1
    tieu de CHUA XU LY (con nguyen mo ta inline tu file mau goc) se dai hon
    dung prefix, khong khop bang so sanh chinh xac."""
    return text in _STOP_EXACT or text.startswith(_STOP_PREFIXES)

_DENSUCCO_LABELS = [f["ten_he_thong"].capitalize() + ": " for f in _DENSUCCO_FORMS]

_CHUA_THE_HIEN = "Chưa thể hiện trên bản vẽ cung cấp."


class BaoCaoThamDinhError(Exception):
    pass


# ---------------------------------------------------------------------------
# Thao tac XML muc thap - dung chung cho toan bo cac ham thay noi dung ben
# duoi (mo rong ky thuat deepcopy da test o cong_van_huong_dan_docx.py sang
# truong hop can XOA/THEM ca doan "mo ta" LAN khoi "Kien nghi:").
# ---------------------------------------------------------------------------
def _text_of_element(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _set_paragraph_text(p, text):
    """Xoa noi dung cac run cu, giu lai run DAU TIEN (giu nguyen dinh dang/
    font) de dat text moi vao - neu doan rong (khong run nao) thi them run
    moi. Dung cung ky thuat voi docx_mergefield.replace_mergefield()."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra.text = ""


def _find_heading_paragraph(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def _capture_kien_nghi_template(doc):
    """Bat 1 cap (doan 'Kien nghi:', doan bullet dau tien ngay sau no) TU FILE
    GOC - PHAI goi 1 LAN DUY NHAT o dau, TRUOC KHI sua bat ky muc nao, de lam
    khuon deepcopy cho cac muc can TU CHEN MOI khoi 'Kien nghi:' (vd muc 8/11
    khi du lieu that co kien nghi nhung mau dang la "khong."). Neu bat truoc
    ma phu thuoc thu tu xu ly cac muc khac se co rui ro: muc dinh muon lam
    nguon co the da bi XOA khoi 'Kien nghi:' cua no vi du lieu that cua CHINH
    muc do rong."""
    for p in doc.paragraphs:
        if p.text.strip() == _KIEN_NGHI_LABEL:
            kn_el = p._element
            next_el = kn_el.getnext()
            if next_el is not None and _text_of_element(next_el).startswith("- "):
                return copy.deepcopy(kn_el), copy.deepcopy(next_el)
    return None, None


def _chen_bullet_moi(kien_nghi_p_element, old_bullet_elements, cau_list):
    """Xoa cac bullet CU ngay sau kien_nghi_p_element, chen bullet MOI cho
    tung cau trong cau_list (deepcopy 1 bullet mau de giu dinh dang) - giong
    het ky thuat _thay_danh_sach_kien_nghi() cua cong_van_huong_dan_docx.py."""
    template_bullet = old_bullet_elements[0] if old_bullet_elements else kien_nghi_p_element
    anchor = old_bullet_elements[-1] if old_bullet_elements else kien_nghi_p_element

    for cau in cau_list:
        new_el = copy.deepcopy(template_bullet)
        for t in new_el.iter(qn("w:t")):
            t.text = ""
        first_t = next(new_el.iter(qn("w:t")), None)
        if first_t is not None:
            first_t.text = f"- {cau}"
        anchor.addnext(new_el)
        anchor = new_el

    for el in old_bullet_elements:
        el.getparent().remove(el)


def _chen_khoi_kien_nghi_moi(heading_p, cau_list, kien_nghi_template_pair):
    """Chen MOI 1 khoi 'Kien nghi:' + bullet ngay sau heading_p, muon dinh
    dang tu kien_nghi_template_pair (deepcopy tu file GOC, xem
    _capture_kien_nghi_template) - dung cho muc dang KHONG co san khoi 'Kien
    nghi:' trong mau (vd muc 8/11) nhung du lieu THAT cua du an nay co kien
    nghi can dien."""
    kn_template, bullet_template = kien_nghi_template_pair
    if kn_template is None:
        return  # mau khong co khoi "Kien nghi:" nao de muon dinh dang - bo qua an toan
    new_kn = copy.deepcopy(kn_template)
    heading_p._element.addnext(new_kn)
    anchor = new_kn
    for cau in cau_list:
        new_el = copy.deepcopy(bullet_template)
        for t in new_el.iter(qn("w:t")):
            t.text = ""
        first_t = next(new_el.iter(qn("w:t")), None)
        if first_t is not None:
            first_t.text = f"- {cau}"
        anchor.addnext(new_el)
        anchor = new_el


def _thay_heading_va_kien_nghi(doc, prefix, mo_ta, kien_nghi_list, kien_nghi_template_pair):
    """Ap dung cho 10/11 muc (tru muc 2). Tim doan co prefix, GOP toan bo
    (cac) doan "mo ta" phu (neu co, xem module-docstring) vao dung 1 doan
    "prefix + mo_ta", roi xu ly khoi "Kien nghi:" theo dung 3 truong hop:
    (a) co san + du lieu that co kien nghi -> thay bullet; (b) co san + du
    lieu that KHONG co kien nghi -> xoa het khoi; (c) KHONG co san + du lieu
    that CO kien nghi -> tu chen moi khoi (muon dinh dang tu noi khac)."""
    heading_p = _find_heading_paragraph(doc, prefix)
    if heading_p is None:
        return  # mau khong co dung tieu de nay - bo qua, khong loi (phong mau doi)

    body = heading_p._element.getparent()
    all_ps = list(body.findall(qn("w:p")))
    start_idx = all_ps.index(heading_p._element)

    extra_mota_elements = []
    idx = start_idx + 1
    while idx < len(all_ps):
        el = all_ps[idx]
        text = _text_of_element(el)
        if _la_diem_dung(text) or text.startswith("- ") or text == "":
            break
        extra_mota_elements.append(el)
        idx += 1

    _set_paragraph_text(heading_p, f"{prefix} {mo_ta}".rstrip())
    for el in extra_mota_elements:
        el.getparent().remove(el)

    # Doan NGAY SAU heading_p (theo XML THAT SU sau khi da xoa cac doan mo ta
    # phu o tren, KHONG phai list all_ps cu) moi la ung vien "Kien nghi:".
    next_el = heading_p._element.getnext()
    next_text = _text_of_element(next_el) if next_el is not None else None

    if next_text == _KIEN_NGHI_LABEL:
        kn_all_ps = list(body.findall(qn("w:p")))
        kn_start_idx = kn_all_ps.index(next_el)
        old_bullets = []
        j = kn_start_idx + 1
        while j < len(kn_all_ps):
            el = kn_all_ps[j]
            text = _text_of_element(el)
            if text.startswith("- ") or text == "":
                old_bullets.append(el)
                j += 1
            else:
                break
        if not kien_nghi_list:
            for el in old_bullets:
                el.getparent().remove(el)
            next_el.getparent().remove(next_el)
        else:
            _chen_bullet_moi(next_el, old_bullets, kien_nghi_list)
    elif kien_nghi_list:
        _chen_khoi_kien_nghi_moi(heading_p, kien_nghi_list, kien_nghi_template_pair)


def _thay_heading_bullet_truc_tiep(doc, prefix, danh_sach_dong):
    """Rieng cho muc 2 - cau truc "heading -> bullet THANG" (giong VBHD),
    KHONG co doan "Kien nghi:" xen giua. GIU NGUYEN doan tieu de (khong doi
    text), chi thay cac doan "- ..." ngay sau bang danh_sach_dong."""
    heading_p = _find_heading_paragraph(doc, prefix)
    if heading_p is None:
        return

    body = heading_p._element.getparent()
    all_ps = list(body.findall(qn("w:p")))
    start_idx = all_ps.index(heading_p._element)

    old_bullets = []
    idx = start_idx + 1
    while idx < len(all_ps):
        el = all_ps[idx]
        text = _text_of_element(el)
        if text.startswith("- ") or text == "":
            old_bullets.append(el)
            idx += 1
        else:
            break

    if not old_bullets:
        return  # khong co bullet mau nao de muon dinh dang - bo qua an toan

    if not danh_sach_dong:
        for el in old_bullets:
            el.getparent().remove(el)
        return

    template_bullet = old_bullets[0]
    anchor = old_bullets[-1]
    for dong in danh_sach_dong:
        text = dong if dong.startswith("- ") else f"- {dong}"
        new_el = copy.deepcopy(template_bullet)
        for t in new_el.iter(qn("w:t")):
            t.text = ""
        first_t = next(new_el.iter(qn("w:t")), None)
        if first_t is not None:
            first_t.text = text
        anchor.addnext(new_el)
        anchor = new_el

    for el in old_bullets:
        el.getparent().remove(el)


# ---------------------------------------------------------------------------
# Tong hop noi dung 11 muc tu quy_mo + hang_muc_list
# ---------------------------------------------------------------------------
def _flatten_kien_nghi(kien_nghi: dict) -> list:
    out = []
    for key in ("I_chua_the_hien", "II_chua_thong_nhat", "III_chua_phu_hop", "IV_de_xuat_bo_sung"):
        out.extend(kien_nghi.get(key) or [])
    return out


def _tach_tong_ket_theo_nhan(tong_ket_gop: str, nhan: str) -> str:
    """densucco tra ve tong_ket gop dang '<Nhan1>: cau1. <Nhan2>: cau2.' (xem
    densucco_reader.read_drawing(), bien tong_ket_parts, cac phan noi bang
    dung 1 dau cach) - tach best-effort phan ung voi 1 nhan cu the: tim
    marker '{nhan.capitalize()}: ', cat toi ngay TRUOC marker cua nhan KHAC
    (biet truoc CA 2 nhan tu chinh densucco_reader.FORMS, khong doan/dung
    regex long leo) hoac het chuoi neu khong co nhan nao khac theo sau. Tra
    ve rong neu khong tim thay marker cua chinh nhan can tach."""
    if not tong_ket_gop:
        return ""
    marker = nhan.capitalize() + ": "
    idx = tong_ket_gop.find(marker)
    if idx == -1:
        return ""
    start = idx + len(marker)
    end = len(tong_ket_gop)
    for other_marker in _DENSUCCO_LABELS:
        if other_marker == marker:
            continue
        other_idx = tong_ket_gop.find(other_marker, start)
        if other_idx != -1:
            end = min(end, other_idx)
    return tong_ket_gop[start:end].strip()


def _mo_ta_quy_mo(quy_mo: dict) -> str:
    """1 cau van phong ngan gon mo ta quy mo cong trinh cho muc 4, CHI dung
    field THAT SU co trong quy_mo (bo qua field nao la None, khong tu bia so
    0) - tai dung _fmt_so() (dinh dang so kieu Viet Nam) tu quy_mo_store.py,
    khong viet trung logic dinh dang."""
    parts = []
    floors = quy_mo.get("floors")
    basements = quy_mo.get("basements")
    semi = quy_mo.get("semiBasements")
    if floors is not None or basements or semi:
        so_tang_parts = []
        if basements:
            so_tang_parts.append(f"{_fmt_so(basements)} tầng hầm")
        if semi:
            so_tang_parts.append(f"{_fmt_so(semi)} tầng bán hầm")
        if floors is not None:
            so_tang_parts.append(f"{_fmt_so(floors)} tầng nổi")
        if so_tang_parts:
            parts.append("Số tầng: " + " và ".join(so_tang_parts) + ".")
    if quy_mo.get("totalArea") is not None:
        parts.append(f"Tổng diện tích sàn: {_fmt_so(quy_mo['totalArea'])} m².")
    if quy_mo.get("volume") is not None:
        parts.append(f"Tổng khối tích: {_fmt_so(quy_mo['volume'])} m³.")
    if quy_mo.get("hFire") is not None:
        parts.append(f"Chiều cao PCCC: {_fmt_so(quy_mo['hFire'])} m.")
    return " ".join(parts) if parts else "Chưa xác định — cần bổ sung dữ liệu quy mô."


def _noi_dung_muc(quy_mo: dict, hang_muc_list: list) -> dict:
    by_slot = {h.get("slot"): h for h in hang_muc_list}
    out = {}

    # Muc 1 — RULE-BASED (evaluate_tham_dinh), KHONG goi AI. Cau 1 lay tu
    # quy_mo (khong xac dinh lich su tham duyet - ghi ro caveat). Cau 2 =
    # dung nguyen r["detail"] cua evaluate_tham_dinh (da la cau van phong
    # phap ly san co, dung y het cho Form A id=2 - xem quy_mo_store.py).
    try:
        r = evaluate_tham_dinh(quy_mo) if quy_mo.get("occ") else None
    except ThamDinhInputError:
        r = None
    ten = quy_mo.get("tenCongTrinh") or "công trình nêu trên"
    mo_ta_1 = (
        f"Qua tra cứu dữ liệu lưu trữ, {ten} — CHƯA XÁC ĐỊNH được lịch sử thẩm duyệt/nghiệm thu PCCC "
        f"trước đây (cần cán bộ tự kiểm tra). "
        + (r["detail"] if r else "Chưa đủ dữ liệu quy mô để xác định diện thẩm định — cần bổ sung.")
    )
    out[1] = {"mo_ta": mo_ta_1, "kien_nghi": []}

    # Muc 2 — thang tu thanhPhanHoSo (Buoc 2.5), KHONG goi AI.
    out[2] = {"mo_ta": "", "kien_nghi": list(quy_mo.get("thanhPhanHoSo") or [])}

    # Muc 3 — rut tu quy_mo, KHONG can hang_muc nao.
    parts3 = []
    if quy_mo.get("hazard"):
        parts3.append(f"hạng nguy hiểm cháy nổ {quy_mo['hazard']}")
    if quy_mo.get("bacChiuLua"):
        parts3.append(f"bậc chịu lửa {quy_mo['bacChiuLua']}")
    if quy_mo.get("capNguyHiemChayKetCau"):
        parts3.append(f"cấp nguy hiểm cháy kết cấu {quy_mo['capNguyHiemChayKetCau']}")
    out[3] = {"mo_ta": (", ".join(parts3) + ".") if parts3 else "Chưa xác định — cần bổ sung.", "kien_nghi": []}

    # Muc 4 — mo ta rut tu quy_mo, kien nghi = canh bao quy mo (slot "quy_mo"
    # trong hang_muc_list, dung dung tinh than VBHD nhom "Thong tin cong trinh").
    mota4 = _mo_ta_quy_mo(quy_mo)
    kn4 = (by_slot.get("quy_mo") or {}).get("kien_nghi") or {}
    out[4] = {"mo_ta": mota4, "kien_nghi": _flatten_kien_nghi(kn4)}

    # Muc 5 — densucco GOP (den+loa+binh), dat CA kien_nghi vao day (xem
    # module-docstring cong_van_huong_dan_docx.py / prompt Pha 3 Buoc 4 cho
    # ly do: densucco_reader.read_drawing() KHONG tach duoc kien_nghi theo
    # tung sub-form, chi tach duoc items).
    densucco = by_slot.get("densucco") or {}
    mota5 = _tach_tong_ket_theo_nhan(densucco.get("tong_ket") or "", "Đèn chiếu sáng sự cố, đèn chỉ dẫn thoát nạn")
    out[5] = {"mo_ta": mota5 or _CHUA_THE_HIEN, "kien_nghi": _flatten_kien_nghi(densucco.get("kien_nghi") or {})}

    # Muc 6 — baochay rieng, KHONG co loa (xem ly do tren).
    baochay = by_slot.get("baochay") or {}
    mota6 = (baochay.get("tong_ket") or _CHUA_THE_HIEN) + " (Hệ thống loa thông báo và hướng dẫn thoát nạn — xem đối chiếu gộp tại mục 5.)"
    out[6] = {"mo_ta": mota6, "kien_nghi": _flatten_kien_nghi(baochay.get("kien_nghi") or {})}

    # Muc 7 — ccnuoc.
    ccnuoc = by_slot.get("ccnuoc") or {}
    out[7] = {"mo_ta": ccnuoc.get("tong_ket") or _CHUA_THE_HIEN, "kien_nghi": _flatten_kien_nghi(ccnuoc.get("kien_nghi") or {})}

    # Muc 8 — khibot + botcodinh + giakehang + botchuachay gop (neu co bat ky slot nao).
    slots8 = [by_slot.get(s) for s in ("khibot", "botcodinh", "giakehang", "botchuachay") if by_slot.get(s)]
    if slots8:
        mota8 = " ".join((s.get("tong_ket") or "") for s in slots8).strip() or _CHUA_THE_HIEN
        kn8 = []
        for s in slots8:
            kn8.extend(_flatten_kien_nghi(s.get("kien_nghi") or {}))
        out[8] = {"mo_ta": mota8, "kien_nghi": kn8}
    else:
        out[8] = {"mo_ta": "không.", "kien_nghi": []}

    # Muc 9 — phan binh trong densucco (khong tach duoc kien nghi, xem tren).
    mota9 = _tach_tong_ket_theo_nhan(densucco.get("tong_ket") or "", "Bình chữa cháy xách tay/xe đẩy")
    if mota9:
        mota9 += " (Kiến nghị liên quan bình chữa cháy xách tay — xem mục 5.)"
    out[9] = {"mo_ta": mota9 or "không.", "kien_nghi": []}

    # Muc 10 — dienpccc.
    dien = by_slot.get("dienpccc") or {}
    out[10] = {"mo_ta": dien.get("tong_ket") or _CHUA_THE_HIEN, "kien_nghi": _flatten_kien_nghi(dien.get("kien_nghi") or {})}

    # Muc 11 — luon "khong." (chua co nguon du lieu nao khac trong app).
    out[11] = {"mo_ta": "không.", "kien_nghi": []}

    return out


def build_bao_cao_tham_dinh_docx(session_data: dict, hang_muc_list: list) -> bytes:
    """session_data: {"quy_mo": dict tu quy_mo_store.get_quy_mo() (co the
    None)}. hang_muc_list: list dict {slot, tong_ket, kien_nghi} - CHU Y them
    "tong_ket" so voi cong_van_huong_dan_docx.py (buoc do khong can tong_ket,
    buoc nay can de dung cho phan mo ta hien trang cua tung muc)."""
    if not TEMPLATE_PATH.exists():
        raise BaoCaoThamDinhError("Chưa có file mẫu báo cáo thẩm định — liên hệ quản trị hệ thống.")

    doc = Document(str(TEMPLATE_PATH))
    quy_mo = session_data.get("quy_mo") or {}

    replace_mergefield(doc, "tên_công_trình", quy_mo.get("tenCongTrinh"))
    replace_mergefield(doc, "chủ_đầu_tư", quy_mo.get("chuDauTu"))
    replace_mergefield(doc, "ĐỊA_ĐIỂM_XÂY_DỰNG", quy_mo.get("diaDiemXayDung"))
    replace_mergefield(doc, "địa_chỉ_chủ_đầu_tư", quy_mo.get("diaChiChuDauTu"))
    replace_mergefield(doc, "ĐƠN_VỊ_TƯ_VẤN_THIẾT_KẾ_PCCC", quy_mo.get("donViTuVanThietKe"))
    replace_mergefield(doc, "MÃ_HỒ_SƠ", quy_mo.get("maHoSo"))
    replace_mergefield(doc, "TỔNG_MỨC_ĐẦU_TƯ_XÂY_DỰNG", quy_mo.get("tongMucDauTu"))
    replace_mergefield(doc, "CÁN_BỘ", None)  # ten nguoi - luon do, khong co nguon du lieu
    replace_mergefield(doc, "GIẤY_CHỨNG_NHẬN_QUYỀN_SỬ_DỤNG_ĐẤT", None)  # chua co field rieng, luon do

    # Bat khuon "Kien nghi:" + bullet TU FILE GOC 1 LAN DUY NHAT, TRUOC KHI
    # sua bat ky muc nao (xem docstring _capture_kien_nghi_template).
    kien_nghi_template_pair = _capture_kien_nghi_template(doc)

    noi_dung = _noi_dung_muc(quy_mo, hang_muc_list)

    for so_muc, prefix in _MUC_PREFIX.items():
        info = noi_dung.get(so_muc) or {"mo_ta": "", "kien_nghi": []}
        if so_muc in _MUC_BULLET_TRUC_TIEP:
            _thay_heading_bullet_truc_tiep(doc, prefix, info["kien_nghi"])
        else:
            _thay_heading_va_kien_nghi(doc, prefix, info["mo_ta"], info["kien_nghi"], kien_nghi_template_pair)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
