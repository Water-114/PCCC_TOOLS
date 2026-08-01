"""Dựng file .docx tổng hợp kiến nghị thiết kế từ danh sách hạng mục đã đọc AI
(báo cháy, điện PCCC, chữa cháy nước, ...) — KHÔNG dùng template, dựng mới hoàn
toàn bằng python-docx. Mỗi hạng mục 1 khối "KIẾN NGHỊ THIẾT KẾ - {tên hệ thống}",
gộp nhiều hạng mục vào cùng 1 file.
"""

import io

from docx import Document

KIEN_NGHI_NHOM = [
    ("I_chua_the_hien", "I. Nội dung chưa thể hiện"),
    ("II_chua_thong_nhat", "II. Nội dung chưa thống nhất"),
    ("III_chua_phu_hop", "III. Nội dung chưa phù hợp QCVN, TCVN"),
    ("IV_de_xuat_bo_sung", "IV. Đề xuất bổ sung hồ sơ"),
]

FILENAME = "Kien_nghi_thiet_ke.docx"


def build_kien_nghi_docx(hang_muc_list: list) -> bytes:
    """hang_muc_list: list dict {ten_he_thong, so_hieu_ban_ve, kien_nghi}, trong đó
    kien_nghi là dict 4 khoá (I_chua_the_hien..IV_de_xuat_bo_sung) -> list câu. Trả
    về bytes nội dung file .docx gộp tất cả hạng mục."""
    doc = Document()

    for idx, hang_muc in enumerate(hang_muc_list):
        if idx > 0:
            doc.add_page_break()

        ten_he_thong = (hang_muc.get("ten_he_thong") or "").strip()
        so_hieu = (hang_muc.get("so_hieu_ban_ve") or "").strip() or "Không xác định được số hiệu bản vẽ"
        kien_nghi = hang_muc.get("kien_nghi") or {}

        doc.add_heading(f"KIẾN NGHỊ THIẾT KẾ - {ten_he_thong}", level=1)

        so_hieu_p = doc.add_paragraph()
        so_hieu_run = so_hieu_p.add_run(f"Bản vẽ số: {so_hieu}")
        so_hieu_run.italic = True

        for key, label in KIEN_NGHI_NHOM:
            doc.add_heading(label, level=2)
            items = kien_nghi.get(key) or []
            if items:
                for item in items:
                    doc.add_paragraph(str(item), style="List Number")
            else:
                doc.add_paragraph("(Không có)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
