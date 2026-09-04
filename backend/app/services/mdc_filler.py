"""Trích tiêu chí thật từ mẫu MĐC (bảng đối chiếu) và điền lại cột "Nội dung
thiết kế" + "Kết luận" theo kết quả AI đối chiếu bản vẽ. Dùng chung cho mọi
hạng mục (báo cháy B1/B2, điện PCCC B14, ...) — mỗi hạng mục chỉ cần thêm 1
khoá vào TEMPLATE_PATHS/TEMPLATE_FILENAMES bên dưới.

Nguyên tắc theo đúng references/quy-tac-dien-form.md của skill ra-mau-doi-chieu-pccc:
chỉ điền cột 3, tô đỏ nội dung mới, giữ nguyên 100% các cột còn lại và bố cục bảng.
Cột "Kết luận" (cột 5): tô đỏ CHỈ khi giá trị là "KN" (kiến nghị/chưa đạt) —
"Đạt" và rỗng (mục không áp dụng) giữ màu mặc định. Áp dụng chung cho MỌI
hạng mục, không riêng form nào.
"""

import io
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

TEMPLATES_DIR = Path(__file__).resolve().parent / "mdc_templates"
TEMPLATE_PATHS = {
    "thuong": TEMPLATES_DIR / "B1_bao_chay_thuong.docx",
    "dia_chi": TEMPLATES_DIR / "B2_bao_chay_dia_chi.docx",
    "dien_pccc": TEMPLATES_DIR / "B14_dien_pccc.docx",
    "tram_bom": TEMPLATES_DIR / "B3_tram_bom.docx",
    "hong_nuoc": TEMPLATES_DIR / "B5_hong_nuoc.docx",
    "chua_chay_tu_dong": TEMPLATES_DIR / "B6_chua_chay_tu_dong.docx",
    "binh_chua_chay": TEMPLATES_DIR / "B12_binh_chua_chay.docx",
    "den_su_co": TEMPLATES_DIR / "B13_den_su_co.docx",
    "quy_mo": TEMPLATES_DIR / "A_quy_mo.docx",
    "bot_co_dinh": TEMPLATES_DIR / "B7_bot_co_dinh.docx",
    "khi_hoa_long": TEMPLATES_DIR / "B8_khi_hoa_long.docx",
    "khi_nen": TEMPLATES_DIR / "B9_khi_nen.docx",
    "khi_co2": TEMPLATES_DIR / "B10_khi_co2.docx",
    "sol_khi": TEMPLATES_DIR / "B11_sol_khi.docx",
    "chua_chay_gia_ke_hang": TEMPLATES_DIR / "B15_chua_chay_gia_ke_hang.docx",
    "bot_chua_chay": TEMPLATES_DIR / "B16_bot_chua_chay.docx",
    "A14": TEMPLATES_DIR / "A14_nha_san_xuat.docx",
    "A15": TEMPLATES_DIR / "A15_nha_kho.docx",
}
TEMPLATE_FILENAMES = {
    "thuong": "B1_MDC_bao_chay_thuong.docx",
    "dia_chi": "B2_MDC_bao_chay_dia_chi.docx",
    "dien_pccc": "B14_MDC_dien_pccc.docx",
    "tram_bom": "B3_MDC_tram_bom.docx",
    "hong_nuoc": "B5_MDC_hong_nuoc.docx",
    "chua_chay_tu_dong": "B6_MDC_chua_chay_tu_dong.docx",
    "binh_chua_chay": "B12_MDC_binh_chua_chay.docx",
    "den_su_co": "B13_MDC_den_su_co.docx",
    "quy_mo": "A_MDC_quy_mo.docx",
    "bot_co_dinh": "B7_MDC_bot_co_dinh.docx",
    "khi_hoa_long": "B8_MDC_khi_hoa_long.docx",
    "khi_nen": "B9_MDC_khi_nen.docx",
    "khi_co2": "B10_MDC_khi_co2.docx",
    "sol_khi": "B11_MDC_sol_khi.docx",
    "chua_chay_gia_ke_hang": "B15_MDC_chua_chay_gia_ke_hang.docx",
    "bot_chua_chay": "B16_MDC_bot_chua_chay.docx",
    "A14": "A14_MDC_nha_san_xuat.docx",
    "A15": "A15_MDC_nha_kho.docx",
}

COL_DOI_CHIEU = 1
COL_THIET_KE = 2
COL_QUY_DINH = 3
COL_KHOAN_DIEU = 4
COL_KET_LUAN = 5

MAU_DO = RGBColor(0xEE, 0x00, 0x00)

# Chuyen doi gia tri ket_luan THAT cua AI (dat/chua_dat/chua_the_hien/
# khong_ap_dung - xem ai_schema.KetLuan) sang chu hien thi trong cot 5 cua
# file .docx: "dat" -> "Đạt", "khong_ap_dung" -> rong (muc tuy chon khong
# thiet ke, KHONG phai loi), moi gia tri khac (chua_dat/chua_the_hien) ->
# "KN" (kien nghi). Chuyen ra day (thay vi de rieng trong routes/aiho.py) de
# la 1 NGUON DUY NHAT dung chung cho ca 9 route doc tung hang muc (qua
# _answers_from_items()) LAN Form A nguoi dung tu dinh (form_a_upload.py,
# Batch 5A Pha 3 Buoc 5) - tranh 2 noi tu dinh nghia lech nhau theo thoi gian.
KET_LUAN_TO_DOCX = {"dat": "Đạt", "khong_ap_dung": ""}

_ROWS_CACHE = {}


def _extract_rows_from_doc(doc) -> list:
    """Phan LOI cua _extract_rows() cu - nhan thang 1 Document object da mo
    san (khong quan tam Document do mo tu path hay tu BytesIO), tra ve list
    dict {id, doi_chieu, quy_dinh, khoan_dieu}. Dung chung cho ca 18 loai co
    san (qua load_criteria_rows) VA Form A nguoi dung tu dinh (Batch 5A Pha 3
    Buoc 5, xem form_a_upload.py)."""
    table = doc.tables[0]
    rows = []
    for idx, row in enumerate(table.rows):
        if idx == 0:
            continue  # dòng tiêu đề bảng
        doi_chieu = row.cells[COL_DOI_CHIEU].text.strip()
        quy_dinh = row.cells[COL_QUY_DINH].text.strip()
        khoan_dieu = row.cells[COL_KHOAN_DIEU].text.strip()
        if not quy_dinh:
            continue  # dòng "mục" tiêu đề gộp nhóm — không phải dòng tiêu chí thật
        if doi_chieu == quy_dinh == khoan_dieu:
            # Dong tieu de bi merge NGUYEN CA HANG (ca 6 cot) thay vi chi merge
            # dung 1 cot nhan de - phat hien qua thuc te (B12 dong TT=1
            # "Binh chua chay xach tay" bi merge het, khien cot quy_dinh khong
            # rong nhu ky vong). Dau hieu: ca 3 cot deu giong het nhau tung chu -
            # khong the la 1 tieu chi that (quy_dinh/khoan_dieu luon khac
            # doi_chieu trong moi dong tieu chi hop le).
            continue
        rows.append({
            "id": idx,
            "doi_chieu": doi_chieu,
            "quy_dinh": quy_dinh,
            "khoan_dieu": khoan_dieu,
        })
    return rows


def _extract_rows(path):
    return _extract_rows_from_doc(Document(path))


def load_criteria_rows(loai: str) -> list:
    """loai: 'thuong' hoặc 'dia_chi'. Trả về list dict {id, doi_chieu, quy_dinh, khoan_dieu}."""
    if loai not in _ROWS_CACHE:
        _ROWS_CACHE[loai] = _extract_rows(TEMPLATE_PATHS[loai])
    return _ROWS_CACHE[loai]


def fill_doc_in_place(doc, answers: list):
    """Phan LOI cua fill_docx() cu - nhan Document object, DIEN THANG vao do
    (khong tra ve bytes, khong tu doc.save() - de caller tu quyet dinh save
    vao dau). fill_docx() cu goi lai ham nay de khong trung logic."""
    table = doc.tables[0]
    answers_by_id = {a["id"]: a for a in answers if "id" in a}

    for idx, row in enumerate(table.rows):
        if idx == 0:
            continue
        ans = answers_by_id.get(idx)
        if not ans:
            continue
        cell = row.cells[COL_THIET_KE]
        cell.text = ans.get("noi_dung_thiet_ke") or ""
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.color.rgb = MAU_DO

        ket_luan_cell = row.cells[COL_KET_LUAN]
        ket_luan_text = ans.get("ket_luan") or ""
        ket_luan_cell.text = ket_luan_text
        if ket_luan_text == "KN" and ket_luan_cell.paragraphs[0].runs:
            ket_luan_cell.paragraphs[0].runs[0].font.color.rgb = MAU_DO


def fill_docx(loai: str, answers: list) -> bytes:
    """answers: list dict {id, noi_dung_thiet_ke, ket_luan}. Trả về bytes file .docx đã điền."""
    doc = Document(TEMPLATE_PATHS[loai])
    fill_doc_in_place(doc, answers)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def filename_for(loai: str) -> str:
    return TEMPLATE_FILENAMES.get(loai, "MDC_bao_chay.docx")
