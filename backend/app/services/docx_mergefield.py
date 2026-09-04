"""Thay gia tri hien thi cua MERGEFIELD (Word mail-merge field) trong file
.docx, dung chung cho cong van huong dan + bao cao tham dinh (Batch 5A Pha 3).
Chi thay phan "cached result" (run giua fldChar separate/end), GIU NGUYEN
field code (instrText) - Word van mo file binh thuong, khong bao "sua chua
tai lieu". Da test doc-lap voi file mau that truoc khi dua vao code chinh."""

from docx.oxml.ns import qn

from .mdc_filler import MAU_DO

YEU_CAU_NHAP_THONG_TIN = "[yêu cầu nhập thông tin]"


def _iter_all_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_mergefield(document, field_name: str, value):
    """Thay TOAN BO cac cho xuat hien MERGEFIELD "field_name" trong document
    bang value. Neu value la None/rong: dien YEU_CAU_NHAP_THONG_TIN, TO DO
    (mau_do, dung chung voi mdc_filler) de nguoi dung biet can tu nhap. Tra
    ve so luong vi tri da thay (0 neu khong tim thay field nao ten nay)."""
    target_instr = f"MERGEFIELD {field_name}"
    use_value = value if value else YEU_CAU_NHAP_THONG_TIN
    replaced = 0
    for p in _iter_all_paragraphs(document):
        runs = p.runs
        for i, r in enumerate(runs):
            is_target = any(
                it.text and it.text.strip() == target_instr
                for it in r._element.findall(qn("w:instrText"))
            )
            if not is_target:
                continue
            sep_idx = None
            for j in range(i, len(runs)):
                if any(fc.get(qn("w:fldCharType")) == "separate" for fc in runs[j]._element.findall(qn("w:fldChar"))):
                    sep_idx = j
                    break
            if sep_idx is None:
                continue
            end_idx = None
            for k in range(sep_idx + 1, len(runs)):
                if any(fc.get(qn("w:fldCharType")) == "end" for fc in runs[k]._element.findall(qn("w:fldChar"))):
                    end_idx = k
                    break
            if end_idx is None:
                continue
            cached_runs = runs[sep_idx + 1 : end_idx]
            if not cached_runs:
                continue
            cached_runs[0].text = use_value
            if not value:
                cached_runs[0].font.color.rgb = MAU_DO
            for extra in cached_runs[1:]:
                extra.text = ""
            replaced += 1
    return replaced
